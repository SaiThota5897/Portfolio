import streamlit as st
import requests, io, re, os, json
from datetime import datetime
from docx import Document
from typing import Dict, Any, Optional
from streamlit.components.v1 import html as st_html

# ---------------------------
# CONFIG - change if needed
# ---------------------------
GITHUB_OWNER = "SaiThota97"
GITHUB_REPO = "Portfolio"
RESUME_PATH_IN_REPO = "SaiThota_Resume.docx"
BRANCH = "main"  # change if your branch name differs
JSON_CACHE = "resume_data.json"

# Paths to local assets inside the repo
PROFILE_IMG = "assets/profile.jpg"
COMPANY_LOGOS = {
    "Allied World Assurance Company (AWAC)": "assets/logos/awac.jpg",
    "McAfee": "assets/logos/mcafee.jpg",
    "GE HealthCare": "assets/logos/gehealthcare.jpg",
    "N-iX": "assets/logos/nix.jpg"
}
CERT_LOGOS = {
    "Databricks Generative AI Fundamentals": "assets/certs/databricks.png",
    "Microsoft Azure AI Engineer Associate": "assets/certs/azure.png",
    "AWS Cloud Practitioner": "assets/certs/aws.png"
}
EDU_LOGO = "assets/edu/USU_CS.png"

# ---------------------------
# Helpers: GitHub API + raw download
# ---------------------------
def get_github_commits_for_file(owner: str, repo: str, path: str, branch: str = "main", token: Optional[str] = None) -> Optional[dict]:
    url = f"https://api.github.com/repos/{owner}/{repo}/commits"
    params = {"path": path, "sha": branch, "per_page": 1}
    headers = {"Accept": "application/vnd.github.v3+json"}
    if token:
        headers["Authorization"] = f"token {token}"
    r = requests.get(url, params=params, headers=headers, timeout=15)
    if r.status_code == 200:
        data = r.json()
        if isinstance(data, list) and len(data) > 0:
            return data[0]
    return None

def download_raw_file(owner: str, repo: str, path: str, branch: str = "main", token: Optional[str] = None) -> Optional[bytes]:
    raw_url = f"https://raw.githubusercontent.com/{owner}/{repo}/{branch}/{path}"
    headers = {}
    if token:
        headers["Authorization"] = f"token {token}"
    r = requests.get(raw_url, headers=headers, timeout=30)
    if r.status_code == 200:
        return r.content
    return None

# ---------------------------
# Parse .docx
# ---------------------------
def parse_resume_docx_bytes(docx_bytes: bytes) -> Dict[str, Any]:
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    parsed = {
        "name": "",
        "role": "",
        "contact_line": "",
        "summary": "",
        "skills": "",
        "experience": {},
        "certifications": [],
        "education": ""
    }

    if len(paragraphs) >= 1:
        parsed["name"] = paragraphs[0]
    if len(paragraphs) >= 2:
        parsed["role"] = paragraphs[1]
    if len(paragraphs) >= 3:
        parsed["contact_line"] = paragraphs[2]

    section = None
    current_company = None
    company_keys = list(COMPANY_LOGOS.keys())
    for txt in paragraphs[3:]:
        upper = txt.upper()
        if upper.startswith("PROFESSIONAL SUMMARY"):
            section = "summary"
            continue
        if upper.startswith("SKILLS"):
            section = "skills"
            continue
        if upper.startswith("PROFESSIONAL EXPERIENCE"):
            section = "experience"
            continue
        if upper.startswith("CERTIFICATIONS"):
            section = "certifications"
            continue
        if upper.startswith("EDUCATION"):
            section = "education"
            continue

        if section == "experience" and txt.startswith("Client:"):
            remainder = txt[len("Client:"):].strip()
            matched = None
            for c in company_keys:
                if c in remainder:
                    matched = c
                    break
            if matched:
                current_company = matched
                parsed["experience"][current_company] = ""
            else:
                current_company = remainder
                parsed["experience"][current_company] = ""
            continue

        if section == "summary":
            parsed["summary"] += txt + "\n"
        elif section == "skills":
            parsed["skills"] += txt + "\n"
        elif section == "experience":
            if current_company:
                parsed["experience"][current_company] += txt + "\n"
        elif section == "certifications":
            parsed["certifications"].append(txt)
        elif section == "education":
            parsed["education"] += txt + "\n"
        else:
            parsed["summary"] += txt + "\n"

    return parsed
#----------------------------
# Projects parser
#----------------------------
def load_projects_as_markdown_list(docx_path: str) -> str:
    if not os.path.exists(docx_path):
        return "No projects file found."
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    md_list = "\n".join([f"- {para}" for para in paragraphs])  # Markdown bullet points
    return md_list
# ---------------------------
# Load resume from GitHub
# ---------------------------
def load_resume_from_github(owner: str, repo: str, path: str, branch: str = "main", token: Optional[str] = None) -> Dict[str, Any]:
    commit = get_github_commits_for_file(owner, repo, path, branch, token)
    commit_iso = None
    if commit:
        commit_iso = commit.get("commit", {}).get("committer", {}).get("date") or commit.get("commit", {}).get("author", {}).get("date")

    if os.path.exists(JSON_CACHE):
        try:
            with open(JSON_CACHE, "r", encoding="utf-8") as f:
                cached = json.load(f)
            if commit_iso and cached.get("last_updated") == commit_iso:
                return cached["content"]
        except Exception:
            pass

    content_bytes = download_raw_file(owner, repo, path, branch, token)
    if content_bytes is None:
        if os.path.exists(JSON_CACHE):
            with open(JSON_CACHE, "r", encoding="utf-8") as f:
                cached = json.load(f)
            return cached.get("content", {})
        raise RuntimeError("Could not download resume file from GitHub and no cache available.")

    parsed = parse_resume_docx_bytes(content_bytes)
    last_updated = commit_iso or datetime.utcnow().isoformat()
    with open(JSON_CACHE, "w", encoding="utf-8") as f:
        json.dump({"last_updated": last_updated, "content": parsed}, f, indent=2)

    return parsed

# ---------------------------
# Linked and GitHub URLs
# ---------------------------
def make_hyperlinked_contact(contact_text: str) -> str:
    if not contact_text:
        return ""

    # Link emails as before
    email_pattern = re.compile(r'[\w\.-]+@[\w\.-]+\.\w+')
    contact_text = email_pattern.sub(
        lambda m: f'<a href="mailto:{m.group(0)}" style="color:#d14836;">{m.group(0)}</a>',
        contact_text
    )

    # Link the words LinkedIn and GitHub (case-insensitive)
    linkedin_url = "https://www.linkedin.com/in/saithota97/"
    github_url = "https://github.com/SaiThota5897"

    contact_text = re.sub(
        r'\bLinkedIn\b',
        f'<a href="{linkedin_url}" target="_blank" style="color:#d14836;">LinkedIn</a>',
        contact_text,
        flags=re.IGNORECASE
    )

    contact_text = re.sub(
        r'\bGitHub\b',
        f'<a href="{github_url}" target="_blank" style="color:#d14836;">GitHub</a>',
        contact_text,
        flags=re.IGNORECASE
    )

    # Replace pipes with spaced pipes for formatting
    contact_text = contact_text.replace("|", "&nbsp;|&nbsp;")

    return contact_text

# ---------------------------
# Main UI
# ---------------------------
def main():
    st.set_page_config(page_title="Sai Thota — Portfolio(Resume)", layout="wide")

    st.markdown("""
        <style>
        a[id] {
        scroll-margin-top: 150px;
        }
        .streamlit-expanderHeader {
            background-color: #007bff !important;
            color: white !important;
            font-weight: bold !important;
        }
        div[data-testid="stExpanderContent"] {
            background-color: #007bff !important;
            color: white !important;
            padding: 10px;
            border-radius: 6px;
        }
        </style>
    """, unsafe_allow_html=True)

    # Load resume data once
    token = None
    try:
        token = st.secrets.get("GITHUB_TOKEN", None)
    except Exception:
        token = None

    with st.spinner("Loading resume from GitHub..."):
        try:
            resume = load_resume_from_github(GITHUB_OWNER, GITHUB_REPO, RESUME_PATH_IN_REPO, BRANCH, token)
        except Exception as e:
            st.error(f"Error loading resume: {e}")
            st.stop()

    # Initialize or get current page from session state
    if "current_page" not in st.session_state:
        st.session_state["current_page"] = "resume"  # default page

    # CSS and styles for sticky header, buttons
    st.markdown(
        """
        <style>
        div[role="button"] > button {
            background-color: #007bff;
            color: white;
            border-radius: 8px;
            padding: 0.25em 0.8em;
            font-weight: 600;
            box-shadow: 0 2px 5px rgb(0 0 0 / 0.15);
            transition: background-color 0.3s ease;
            cursor: pointer;
        }
        div[role="button"] > button:hover {
            background-color: #0056b3;
        }
        .sticky-header img {
            border-radius: 50%;
            margin-right: 20px;
        }
        .sticky-header-text h1 {
            margin: 0;
            margin-bottom: -34px;  
            color: Gold;
            display: inline-block;
        }
        .sticky-header-text h2 {
            margin: 2px 20px 0 0;
            margin-top: -20px;  
            margin-bottom: -24px;  
            color: LimeGreen;
            display: inline-block;
        }
        .sticky-buttons button {
            background-color: #17a2b8; /* cyan */
            border: none;
            color: white;
            padding: 6px 12px;
            margin-left: 10px;
            border-radius: 4px;
            cursor: pointer;
            font-weight: 600;
            font-size: 14px;
            transition: background-color 0.3s ease;
        }
        .sticky-buttons button:hover {
            background-color: #024f9c;
        }
        .spacer {
            height: 140px;
        }
        .stApp {
            background-color: black;
            color: white;
        }
        /* Ensure widgets and text are readable */
        .stTextInput, .stSelectbox, .stButton > button {
            color: black;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    # Header with profile and navigation buttons (no download)
    profile_img_tag = ""
    if os.path.exists(PROFILE_IMG):
        import base64
        with open(PROFILE_IMG, "rb") as img_file:
            img_b64 = base64.b64encode(img_file.read()).decode()
            profile_img_tag = f"<img src='data:image/jpeg;base64,{img_b64}' width='80'>"

    contact_line = resume.get("contact_line", "")
    contact_html = make_hyperlinked_contact(contact_line)

    st.markdown("""
    <style>
    .stApp {
        background-color: black;
        color: white;
    }
    .stButton > button {
        color: black; !important;
    }
    .sticky-header {
        position: fixed;
        top: 60px;
        left: 0;
        width: 100%;
        background-color: black !important;
        padding: 10px 20px;
        z-index: 50;
        border-bottom: 2px solid #87CEFA; /* light sky blue */
        display: flex;
        align-items: center;
        justify-content: space-between;
        color: white;
    }
    .sticky-header img {
        border-radius: 50%;
        margin-right: 20px;
    }
    .experience-company-name {
    background-color: #87CEEB !important;
    color: black !important;
    padding: 8px 12px;
    border-radius: 4px;
    }
    .sticky-buttons button {
        background-color: #17a2b8; /* cyan */
        border: none;
        color: white;
        padding: 6px 12px;
        margin-left: 10px;
        border-radius: 4px;
        cursor: pointer;
        font-weight: 600;
        font-size: 14px;
        transition: background-color 0.3s ease;
    }
    .sticky-buttons button:hover {
        background-color: #024f9c;
    }
    .spacer {
        height: 140px;
    }
    </style>

    <div class="sticky-header">
        <div style="display:flex; align-items:center;">
            """ + profile_img_tag + """
            <div class="sticky-header-text">
                <h1>""" + resume.get("name", "SAI THOTA") + """</h1>
                <h2>""" + resume.get("role", "AI / Machine Learning Engineer") + """</h2>
                <p>""" + contact_html + """</p>
            </div>
        </div>
        <div class="sticky-buttons">
            <a href="#professional_summary" class="btn-link">Summary</a>
            <a href="#work_experience" class="btn-link">Work Experience</a>
            <a href="#certifications" class="btn-link">Certifications</a>
            <a href="#publication" class="btn-link">Publication</a>
            <a href="#projects" class="btn-link">Projects</a>
            <a href="#education" class="btn-link">Education</a>
            <a href="#about" class="btn-link">About This Page</a>
        </div>
        <style>
            .btn-link {
                background-color: #17a2b8;
                color: white !important;          /* force white text */
                padding: 6px 12px;
                margin-left: 10px;
                border-radius: 4px;
                font-weight: 600;
                font-size: 14px;
                text-decoration: none !important; /* force no underline */
                cursor: pointer;
                display: inline-block;
                transition: background-color 0.3s ease;
            }
            .btn-link:hover {
                background-color: #024f9c;
                text-decoration: none !important; /* force no underline on hover */
            }
            </style>
    </div>
    <div class='spacer'></div>

    """, unsafe_allow_html=True)

    # Use query parameters to detect page
    query_params = st.query_params
    current_page = query_params.get("page", ["resume"])[0]

    # Store current page in session state (optional)
    st.session_state["current_page"] = current_page

    if current_page == "resume":
        # Anchors for scrolling to each section
        st.markdown('<a id="professional_summary"></a>', unsafe_allow_html=True)
        st.header("Professional Summary")
        summary_text = resume.get("summary", "").strip()
        # Bullet point list for summary
        sentences = re.split(r'(?<=[.!?])\s+', summary_text)
        bullets = [s.strip() for s in sentences if s.strip()]
        summary_md = "\n".join([f"- {b}" for b in bullets])
        st.markdown(summary_md)

        st.markdown('<a id="work_experience"></a>', unsafe_allow_html=True)
        st.markdown("---")
        st.subheader("Work Experience")
        companies = list(COMPANY_LOGOS.keys())
        cols = st.columns(len(companies))

        colors = ["#007bff", "#28a745", "#ffc107", "#17a2b8"]  # blue, green, yellow, cyan

        for i, company in enumerate(companies):
            color = colors[i % len(colors)]
            with cols[i]:
                logo_path = COMPANY_LOGOS[company]
                if os.path.exists(logo_path):
                    st.image(logo_path, width=96)
                else:
                    st.write("(logo missing)")

                if st.button(company, key=f"btn_{i}"):
                    st.session_state["selected_company"] = company

                st.markdown(
                    f"""
                    <style>
                    div.stButton > button:focus, div.stButton > button {{
                        background-color: {color} !important;
                        color: white !important;
                        border-radius: 8px !important;
                        font-weight: 600 !important;
                        box-shadow: 0 2px 5px rgb(0 0 0 / 0.15) !important;
                        margin-top: 8px !important;
                        width: 100% !important;
                    }}
                    div.stButton > button:hover {{
                        background-color: {color[:-2]}cc !important;
                    }}
                    </style>
                    """,
                    unsafe_allow_html=True,
                )

        if st.session_state.get("selected_company"):
            st.markdown(
                """
                <style>
                .stExpander > div[data-testid="stExpanderContent"] {
                    background-color: #007bff !important;
                    color: white !important;
                    padding: 10px;
                    border-radius: 6px;
                }
                </style>
                """,
                unsafe_allow_html=True
            )
            comp = st.session_state["selected_company"]
            with st.expander("", expanded=True):
                st.markdown(
                    f'<div class="experience-company-name">{comp} — Experience</div>',
                    unsafe_allow_html=True
                )
                exp_text = resume.get("experience", {}).get(comp, "No details found.")

                # Ensure Role, Description, etc. each start on a new line
                for keyword in ["Role:", "Description:", "Key Contributions:", "Technical Stack:"]:
                    exp_text = exp_text.replace(keyword, f"{keyword}\n")

                # Remove extra spaces from line breaks
                exp_lines = [line.strip() for line in exp_text.split('\n') if line.strip()]

                skip_bullets_prefixes = ("Role:", "Description:", "Key Contributions:", "Technical Stack:")

                exp_md_lines = []
                for line in exp_lines:
                    if any(line.startswith(prefix) for prefix in skip_bullets_prefixes):
                        exp_md_lines.append(f"\n{line}")  # no bullet
                    else:
                        exp_md_lines.append(f"- {line}")  # bullet

                exp_md = "\n".join(exp_md_lines)
                st.markdown(exp_md)

                if st.button("Close (need to click twice)", key="close_expander"):
                    st.session_state["selected_company"] = None

        st.markdown('<a id="certifications"></a>', unsafe_allow_html=True)
        st.markdown("---")
        st.subheader("Certifications")
        certs = list(CERT_LOGOS.keys())
        ccols = st.columns(len(certs))
        for i, cert in enumerate(certs):
            with ccols[i]:
                path = CERT_LOGOS[cert]
                if os.path.exists(path):
                    st.image(path, width=120)
                st.caption(cert)

        st.markdown('<a id="publication"></a>', unsafe_allow_html=True)
        st.markdown("---")
        st.subheader("Publication")
        st.write("I developed advanced hydrological models to assess the impact of climate " \
        "variability on watershed runoff. My research involved integrating remote sensing data and" \
        " climate projections to improve prediction accuracy, supporting sustainable water resource" \
        " management under changing environmental conditions. This work demonstrates my expertise " \
        "in data-driven modeling, environmental analytics, and multidisciplinary collaboration—skills" \
        " highly valuable in data science and AI roles tackling real-world challenges." \
        ""\
        "Technical details can be found at the publication link.")
        st.markdown("[Enhancing Monthly Streamflow Prediction Using Meteorological Factors and" \
        " Machine Learning Models in the Upper Colorado River Basin" \
        "](https://www.mdpi.com/2306-5338/11/5/66)")

        projects_file = "projects.docx"
        if os.path.exists(projects_file):
            projects_md = load_projects_as_markdown_list(projects_file)
            st.markdown('<a id="projects"></a>', unsafe_allow_html=True)
            st.markdown("---")
            st.subheader("Projects")
            st.markdown(projects_md)
        else:
            st.write("Projects file not found.")

        st.markdown('<a id="education"></a>', unsafe_allow_html=True)
        st.markdown("---")
        st.subheader("Education")
        ed_c1, ed_c2 = st.columns([1, 3])
        with ed_c1:
            if os.path.exists(EDU_LOGO):
                st.image(EDU_LOGO, width=120)
        with ed_c2:
            education_text = resume.get("education", "").strip()
            edu_lines = [line.strip() for line in education_text.split('\n') if line.strip()]
            edu_md = "\n".join([f"- {line}" for line in edu_lines])
            st.markdown(edu_md)

        about_file = "aboutpage.txt"
        if os.path.exists(about_file):
            with open(about_file, "r", encoding="utf-8") as f:
                about_text = f.read()
            st.markdown('<a id="about"></a>', unsafe_allow_html=True)  # fixed id typo here
            st.markdown("---")
            st.subheader("About This Page")
            st.markdown(about_text)
        else:
            st.write("About file not found.")

if __name__ == "__main__":
    main()
